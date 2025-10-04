# utils/file_parser.py - 파일 파싱 유틸리티
import io
from typing import Optional, List, Tuple
from utils.internal_vlm import internal_vlm_client


def extract_text_and_images_from_file(file_content: bytes, filename: str) -> Tuple[str, List[bytes]]:
    """
    파일에서 텍스트와 이미지 추출

    Args:
        file_content: 파일 바이너리 내용
        filename: 파일명 (확장자 확인용)

    Returns:
        (추출된 텍스트, 이미지 바이트 리스트)
    """
    filename_lower = filename.lower()
    images = []

    # PDF 파일
    if filename_lower.endswith('.pdf'):
        try:
            import PyPDF2
            import fitz  # PyMuPDF

            # 텍스트 추출 (PyPDF2)
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            text = '\n\n'.join(text_parts)

            # 이미지 추출 (PyMuPDF)
            if internal_vlm_client.is_enabled():
                pdf_doc = fitz.open(stream=file_content, filetype="pdf")
                for page_num in range(pdf_doc.page_count):
                    page = pdf_doc[page_num]
                    image_list = page.get_images(full=True)

                    for img_index, img in enumerate(image_list):
                        xref = img[0]
                        base_image = pdf_doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        images.append(image_bytes)

                pdf_doc.close()

            return text, images

        except ImportError as e:
            if "PyPDF2" in str(e):
                return "[PDF 파싱 라이브러리(PyPDF2)가 설치되지 않았습니다. pip install PyPDF2를 실행하세요.]", []
            elif "fitz" in str(e):
                # fitz가 없으면 이미지 없이 텍스트만 반환
                return text, []
        except Exception as e:
            return f"[PDF 파일 파싱 중 오류 발생: {str(e)}]", []

    # DOCX 파일
    elif filename_lower.endswith('.docx'):
        try:
            from docx import Document
            from docx.oxml.ns import qn

            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)

            # 텍스트 추출
            text_parts = [paragraph.text for paragraph in doc.paragraphs]
            text = '\n\n'.join(text_parts)

            # 이미지 추출
            if internal_vlm_client.is_enabled():
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_data = rel.target_part.blob
                            images.append(image_data)
                        except Exception as img_err:
                            print(f"이미지 추출 중 오류: {img_err}")
                            continue

            return text, images

        except ImportError:
            return "[DOCX 파싱 라이브러리(python-docx)가 설치되지 않았습니다. pip install python-docx를 실행하세요.]", []
        except Exception as e:
            return f"[DOCX 파일 파싱 중 오류 발생: {str(e)}]", []

    # 텍스트 파일
    elif filename_lower.endswith(('.txt', '.md')):
        return file_content.decode('utf-8', errors='ignore'), []

    # DOC 파일
    elif filename_lower.endswith('.doc'):
        return "[DOC 파일은 지원하지 않습니다. DOCX 형식으로 변환해주세요.]", []

    # 기타 파일
    else:
        try:
            return file_content.decode('utf-8', errors='ignore'), []
        except Exception as e:
            return f"[파일 읽기 실패: {str(e)}]", []


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    파일에서 텍스트 추출

    Args:
        file_content: 파일 바이너리 내용
        filename: 파일명 (확장자 확인용)

    Returns:
        추출된 텍스트
    """
    filename_lower = filename.lower()

    # 텍스트 파일
    if filename_lower.endswith(('.txt', '.md')):
        return file_content.decode('utf-8', errors='ignore')

    # PDF 파일
    elif filename_lower.endswith('.pdf'):
        try:
            import PyPDF2
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            return '\n\n'.join(text_parts)
        except ImportError:
            return "[PDF 파싱 라이브러리(PyPDF2)가 설치되지 않았습니다. pip install PyPDF2를 실행하세요.]"
        except Exception as e:
            return f"[PDF 파일 파싱 중 오류 발생: {str(e)}]"

    # DOCX 파일
    elif filename_lower.endswith('.docx'):
        try:
            from docx import Document
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            text_parts = [paragraph.text for paragraph in doc.paragraphs]
            return '\n\n'.join(text_parts)
        except ImportError:
            return "[DOCX 파싱 라이브러리(python-docx)가 설치되지 않았습니다. pip install python-docx를 실행하세요.]"
        except Exception as e:
            return f"[DOCX 파일 파싱 중 오류 발생: {str(e)}]"

    # DOC 파일 (오래된 MS Word 형식)
    elif filename_lower.endswith('.doc'):
        return "[DOC 파일은 지원하지 않습니다. DOCX 형식으로 변환해주세요.]"

    # 기타 파일 - UTF-8로 시도
    else:
        try:
            return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            return f"[파일 읽기 실패: {str(e)}]"
